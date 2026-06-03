from __future__ import annotations

import argparse
import io
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from thesis_render_v2 import ensure_v2_styles, load_style_profile, DEFAULT_STYLE_PROFILE


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FRONT_SOURCE = ROOT / "source" / "front_matter" / "front_matter_source.docx"
DEFAULT_BODY_SOURCE = ROOT / "deliverables" / "docx" / "thesis_render_v2_latest.docx"
DEFAULT_OUTPUT = ROOT / "deliverables" / "docx" / "thesis_render_v2_latest.docx"

def is_chapter1_heading(paragraph) -> bool:
    text = "".join(paragraph.text.split())
    style_name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    if style_name != "heading 1":
        return False
    return text in {"緒論", "第一章緒論"}


def find_chapter1_element(doc: Document):
    for paragraph in doc.paragraphs:
        if is_chapter1_heading(paragraph):
            return paragraph._p
    raise ValueError("Could not find chapter 1 heading in front-matter source DOCX.")


def remove_main_body_from_front_doc(doc: Document) -> None:
    body = doc._element.body
    cutoff = find_chapter1_element(doc)
    removing = False
    for child in list(body.iterchildren()):
        if child is cutoff:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def ensure_main_matter_section_break(doc: Document) -> None:
    """Insert a section break before the first imported chapter heading if missing."""
    body = doc._element.body
    body_sectpr = body.find(qn("w:sectPr"))

    target_para = None
    for paragraph in doc.paragraphs:
        text = "".join(paragraph.text.split())
        style_name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
        if style_name != "heading 1":
            continue
        if text in {"一、緒論", "緒論", "第一章緒論"}:
            target_para = paragraph._element
            break

    if target_para is None:
        return

    siblings = list(body.iterchildren())
    target_idx = siblings.index(target_para)
    if target_idx > 0:
        prev = siblings[target_idx - 1]
        if prev.tag == qn("w:p"):
            prev_ppr = prev.find(qn("w:pPr"))
            if prev_ppr is not None and prev_ppr.find(qn("w:sectPr")) is not None:
                return

    break_para = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    sectpr = deepcopy(body_sectpr) if body_sectpr is not None else OxmlElement("w:sectPr")

    pg_num = sectpr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sectpr.append(pg_num)
    pg_num.set(qn("w:fmt"), "decimal")
    pg_num.set(qn("w:start"), "1")

    ppr.append(sectpr)
    break_para.append(ppr)
    target_para.addprevious(break_para)


def ensure_front_matter_section_breaks(doc: Document, style_profile: dict) -> None:
    policy = style_profile.get("section_policy", {}) if isinstance(style_profile, dict) else {}
    if not isinstance(policy, dict):
        return
    if policy.get("enable_front_matter_section_breaks", True) is False:
        return

    raw_titles = policy.get(
        "front_matter_break_before_titles",
        ["中文摘要", "Abstract", "誌謝", "目錄", "表目錄", "圖目錄"],
    )
    if not isinstance(raw_titles, list):
        return
    target_titles = {"".join(str(item).split()) for item in raw_titles if str(item).strip()}
    if not target_titles:
        return

    # Keep front-matter sectioning, but avoid inserting a trailing break that
    # implicitly carries the new section into Chapter 1.
    avoid_chapter_carryover = policy.get("avoid_chapter_carryover", True) is not False

    body = doc._element.body
    body_sectpr = body.find(qn("w:sectPr"))
    if body_sectpr is None:
        return

    heading_elements = []
    chapter_anchor = None
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
        if style_name != "heading 1":
            continue
        normalized = "".join((paragraph.text or "").split())
        if normalized in {"一、緒論", "緒論", "第一章緒論"} and chapter_anchor is None:
            chapter_anchor = paragraph._element
        if normalized in target_titles:
            heading_elements.append(paragraph._element)

    trailing_front_heading = None
    if avoid_chapter_carryover and chapter_anchor is not None and heading_elements:
        siblings = list(body.iterchildren())
        chapter_idx = siblings.index(chapter_anchor)
        front_candidates = [elem for elem in heading_elements if siblings.index(elem) < chapter_idx]
        if front_candidates:
            trailing_front_heading = max(front_candidates, key=lambda elem: siblings.index(elem))

    for target_para in heading_elements:
        if trailing_front_heading is not None and target_para is trailing_front_heading:
            continue
        siblings = list(body.iterchildren())
        target_idx = siblings.index(target_para)
        if target_idx <= 0:
            continue
        prev = siblings[target_idx - 1]
        if prev.tag == qn("w:p"):
            prev_ppr = prev.find(qn("w:pPr"))
            if prev_ppr is not None and prev_ppr.find(qn("w:sectPr")) is not None:
                continue

        break_para = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        sectpr = deepcopy(body_sectpr)
        ppr.append(sectpr)
        break_para.append(ppr)
        target_para.addprevious(break_para)


def remap_image_relationships(dest_doc: Document, src_doc: Document, element) -> None:
    for blip in element.iter():
        if blip.tag != qn("a:blip"):
            continue
        embed = blip.get(qn("r:embed"))
        if not embed:
            continue
        image_part = src_doc.part.related_parts[embed]
        new_rid, _ = dest_doc.part.get_or_add_image(io.BytesIO(image_part.blob))
        blip.set(qn("r:embed"), new_rid)


def sanitize_section_references(element) -> None:
    for node in element.iter():
        if node.tag != qn("w:sectPr"):
            continue
        for child in list(node):
            if child.tag in {qn("w:headerReference"), qn("w:footerReference")}:
                node.remove(child)


def build_style_id_map(dest_doc: Document, src_doc: Document, style_profile: dict) -> dict[str, str]:
    profile_style_names = list((style_profile.get("styles", {}) or {}).keys())
    style_names = list(dict.fromkeys(profile_style_names + ["Caption", "Caption 圖", "Caption 表"]))
    style_id_map: dict[str, str] = {}
    for style_name in style_names:
        try:
            src_id = src_doc.styles[style_name].style_id
            dest_id = dest_doc.styles[style_name].style_id
        except KeyError:
            continue
        style_id_map[src_id] = dest_id
    return style_id_map


def remap_style_ids(element, style_id_map: dict[str, str]) -> None:
    for node in element.iter():
        if node.tag not in {qn("w:pStyle"), qn("w:rStyle"), qn("w:tblStyle")}:
            continue
        style_id = node.get(qn("w:val"))
        if style_id in style_id_map:
            node.set(qn("w:val"), style_id_map[style_id])


def append_body(dest_doc: Document, src_doc: Document, style_profile: dict) -> None:
    dest_body = dest_doc._element.body
    insert_anchor = list(dest_body.iterchildren())[-1]
    style_id_map = build_style_id_map(dest_doc, src_doc, style_profile)

    for child in list(src_doc._element.body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        clone = deepcopy(child)
        remap_image_relationships(dest_doc, src_doc, clone)
        sanitize_section_references(clone)
        remap_style_ids(clone, style_id_map)
        dest_body.insert(dest_body.index(insert_anchor), clone)


def _make_tof_field_para(doc: Document, style_id: str) -> object:
    """Return a new paragraph element containing a real Word TOF field."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    p = para._element
    # fldChar begin
    r1 = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fc_begin)
    p.append(r1)
    # instrText
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\h \\z \\t "{style_id},1" '
    r2.append(instr)
    p.append(r2)
    # fldChar separate
    r3 = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)
    p.append(r3)
    # placeholder
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "[請在 Word 中更新欄位]"
    r4.append(t)
    p.append(r4)
    # fldChar end
    r5 = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)
    p.append(r5)
    return p


def replace_tof_sections(doc: Document) -> None:
    """Replace hardcoded 'table of figures' paragraphs with real Word TOF fields."""
    paras = doc.paragraphs
    # Find headings and their following hardcoded entries
    targets = [("表目錄", "Caption 表"), ("圖目錄", "Caption 圖")]
    for heading_text, style_id in targets:
        heading_para = None
        for p in paras:
            if p.style and p.style.name.strip().lower() == "heading 1":
                try:
                    if heading_text in p.text:
                        heading_para = p
                        break
                except Exception:
                    pass
        if heading_para is None:
            continue
        # Remove all old generated/static directory lines until next Heading 1.
        body = doc._element.body
        siblings = list(body.iterchildren())
        h_idx = siblings.index(heading_para._element)
        to_remove = []
        for elem in siblings[h_idx + 1:]:
            if elem.tag != qn("w:p"):
                break
            pStyle = elem.find('.//' + qn('w:pStyle'))
            style_val = pStyle.get(qn('w:val'), '') if pStyle is not None else ''
            is_heading1 = style_val in {'11', 'Heading1', 'Heading 1'}
            if is_heading1:
                break
            to_remove.append(elem)
        for elem in to_remove:
            body.remove(elem)
        # Insert real TOF field after heading
        tof_p = _make_tof_field_para(doc, style_id)
        heading_para._element.addnext(tof_p)


def merge(front_source: Path, body_source: Path, output: Path) -> None:
    front_doc = Document(front_source)
    body_doc = Document(body_source)
    style_profile = load_style_profile(DEFAULT_STYLE_PROFILE)

    ensure_v2_styles(front_doc, style_profile=style_profile, preserve_front_matter=True)
    remove_main_body_from_front_doc(front_doc)
    ensure_front_matter_section_breaks(front_doc, style_profile)
    append_body(front_doc, body_doc, style_profile)
    # Disabled by design: chapter-start section breaks are p-whitelist-controlled.
    # ensure_main_matter_section_break(front_doc)
    replace_tof_sections(front_doc)
    ensure_front_matter_section_breaks(front_doc, style_profile)

    output.parent.mkdir(parents=True, exist_ok=True)
    front_doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Merge legacy front matter with v2 body DOCX.")
    parser.add_argument("--front-source", type=Path, default=DEFAULT_FRONT_SOURCE)
    parser.add_argument("--body-source", type=Path, default=DEFAULT_BODY_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    merge(args.front_source.resolve(), args.body_source.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
