from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "assets" / "templates" / "docx" / "thesis_template_thin.docx"


def add_field_run(paragraph, field_code: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    paragraph.add_run("1")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def set_east_asia_font(style, east_asia: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def build_template(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    set_east_asia_font(normal, "標楷體")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(24)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    set_east_asia_font(heading1, "標楷體")
    heading1.paragraph_format.space_before = Pt(9)
    heading1.paragraph_format.space_after = Pt(9)
    heading1.paragraph_format.page_break_before = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(16)
    heading2.font.bold = True
    set_east_asia_font(heading2, "標楷體")
    heading2.paragraph_format.space_before = Pt(9)
    heading2.paragraph_format.space_after = Pt(6)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Times New Roman"
    heading3.font.size = Pt(14)
    heading3.font.bold = True
    set_east_asia_font(heading3, "標楷體")
    heading3.paragraph_format.space_before = Pt(6)
    heading3.paragraph_format.space_after = Pt(4)

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(11)
    caption.font.bold = True
    set_east_asia_font(caption, "標楷體")
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.first_line_indent = Pt(0)

    try:
        note_style = doc.styles["圖表註"]
    except KeyError:
        note_style = doc.styles.add_style("圖表註", WD_STYLE_TYPE.PARAGRAPH)
    note_style.base_style = normal
    note_style.font.name = "Times New Roman"
    note_style.font.size = Pt(10.5)
    set_east_asia_font(note_style, "標楷體")
    note_style.paragraph_format.left_indent = Pt(54)
    note_style.paragraph_format.first_line_indent = Pt(-54)
    note_style.paragraph_format.space_before = Pt(0)
    note_style.paragraph_format.space_after = Pt(4)
    note_style.paragraph_format.line_spacing = 1.0

    try:
        appendix_heading = doc.styles["Appendix Heading"]
    except KeyError:
        appendix_heading = doc.styles.add_style("Appendix Heading", WD_STYLE_TYPE.PARAGRAPH)
    appendix_heading.base_style = doc.styles["Normal"]
    appendix_heading.font.name = "Times New Roman"
    appendix_heading.font.size = Pt(18)
    appendix_heading.font.bold = True
    set_east_asia_font(appendix_heading, "標楷體")
    appendix_heading.paragraph_format.space_before = Pt(9)
    appendix_heading.paragraph_format.space_after = Pt(9)
    appendix_heading.paragraph_format.keep_with_next = True
    appendix_heading.paragraph_format.page_break_before = True
    # outlineLvl=0 so appendix headings appear in TOC at chapter level
    pPr = appendix_heading.element.get_or_add_pPr()
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "0")
    pPr.append(outlineLvl)

    # TOC styles: non-bold, proper indentation
    toc_indents = [Pt(0), Pt(12), Pt(24), Pt(36)]
    for level in range(1, 5):
        sname = f"TOC {level}"
        try:
            toc_style = doc.styles[sname]
        except KeyError:
            toc_style = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
        toc_style.font.bold = False
        toc_style.font.name = "Times New Roman"
        toc_style.font.size = Pt(12)
        set_east_asia_font(toc_style, "標楷體")
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(2)
        toc_style.paragraph_format.left_indent = toc_indents[level - 1]
        toc_style.paragraph_format.first_line_indent = Pt(0)

    # Caption sub-styles for separate 圖目錄 / 表目錄 TOF fields
    for sub in ("圖", "表"):
        sname = f"Caption {sub}"
        try:
            csub = doc.styles[sname]
        except KeyError:
            csub = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
        csub.base_style = doc.styles["Caption"]

    for style_name in ["Thesis Level 4", "Thesis Level 4 Lead", "Thesis Level 4 Body", "Thesis Level 5", "Thesis Level 5 Body"]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        set_east_asia_font(style, "標楷體")

    doc.styles["Thesis Level 4"].font.bold = True
    doc.styles["Thesis Level 4"].paragraph_format.space_before = Pt(3)
    doc.styles["Thesis Level 4"].paragraph_format.space_after = Pt(0)
    doc.styles["Thesis Level 4"].paragraph_format.first_line_indent = Pt(0)

    doc.styles["Thesis Level 4 Lead"].paragraph_format.space_before = Pt(0)
    doc.styles["Thesis Level 4 Lead"].paragraph_format.space_after = Pt(0)
    doc.styles["Thesis Level 4 Lead"].paragraph_format.first_line_indent = Pt(24)

    doc.styles["Thesis Level 4 Body"].paragraph_format.space_before = Pt(0)
    doc.styles["Thesis Level 4 Body"].paragraph_format.space_after = Pt(0)
    doc.styles["Thesis Level 4 Body"].paragraph_format.first_line_indent = Pt(24)

    doc.styles["Thesis Level 5"].font.bold = True
    doc.styles["Thesis Level 5"].paragraph_format.space_before = Pt(3)
    doc.styles["Thesis Level 5"].paragraph_format.space_after = Pt(0)
    doc.styles["Thesis Level 5"].paragraph_format.left_indent = Pt(18)
    doc.styles["Thesis Level 5"].paragraph_format.first_line_indent = Pt(0)

    doc.styles["Thesis Level 5 Body"].paragraph_format.space_before = Pt(0)
    doc.styles["Thesis Level 5 Body"].paragraph_format.space_after = Pt(0)
    doc.styles["Thesis Level 5 Body"].paragraph_format.left_indent = Pt(18)
    doc.styles["Thesis Level 5 Body"].paragraph_format.first_line_indent = Pt(24)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field_run(footer_para, " PAGE ")

    if doc.paragraphs:
        doc.paragraphs[0].clear()
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a thin thesis DOCX template for v2 rendering.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_template(args.output.resolve())


if __name__ == "__main__":
    main()
