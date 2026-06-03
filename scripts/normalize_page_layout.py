from __future__ import annotations

import argparse
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from thesis_shadow import NS, WORD_NAMESPACE_DECLS, write_docx_xml


if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


for prefix, uri in WORD_NAMESPACE_DECLS.items():
    ET.register_namespace(prefix, uri)


PORTRAIT_PAGE = (Cm(21.0), Cm(29.7))
LANDSCAPE_PAGE = (Cm(29.7), Cm(21.0))

PORTRAIT_MARGINS = {
    "top": Cm(4.0),
    "bottom": Cm(3.0),
    "left": Cm(3.0),
    "right": Cm(3.0),
}

LANDSCAPE_MARGINS = {
    "top": Cm(4.0),
    "bottom": Cm(3.0),
    "left": Cm(3.0),
    "right": Cm(3.0),
}

HEADER_DISTANCE = Cm(1.5)
FOOTER_DISTANCE = Cm(2.0)
GUTTER = Cm(0.0)

# Section policy in this thesis template:
#   1. cover/title preliminaries (no numbering restart here)
#   2. front matter, lower roman starting at i
#   3. main matter, arabic starting at 1 (from Chapter 1)
ROMAN_START_MARKERS = {"中文摘要", "摘要"}
CHAPTER_START_MARKERS = {"一、緒論", "緒論", "第一章緒論"}

W = NS["w"]


def paragraph_text(paragraph: ET.Element) -> str:
    parts: list[str] = []
    for node in paragraph.findall(f".//{{{W}}}t"):
        if node.text:
            parts.append(node.text)
    return "".join(parts).strip()


def compact_text(text: str) -> str:
    return "".join(text.split())


def paragraph_style_id(paragraph: ET.Element) -> str:
    ppr = paragraph.find(f"{{{W}}}pPr")
    if ppr is None:
        return ""
    pstyle = ppr.find(f"{{{W}}}pStyle")
    if pstyle is None:
        return ""
    return pstyle.get(f"{{{W}}}val", "")


def ensure_main_matter_section_break(root: ET.Element) -> bool:
    body = root.find(f".//{{{W}}}body")
    if body is None:
        return False

    body_sectpr = body.find(f"{{{W}}}sectPr")
    if body_sectpr is None:
        return False

    children = list(body)
    target_idx = None
    for idx, node in enumerate(children):
        if node.tag != f"{{{W}}}p":
            continue
        text = "".join(paragraph_text(node).split())
        style_id = paragraph_style_id(node)
        if style_id not in {"1", "11", "Heading1", "Heading 1"}:
            continue
        if text in {"一、緒論", "緒論", "第一章緒論"}:
            target_idx = idx
            break

    if target_idx is None or target_idx == 0:
        return False

    prev = children[target_idx - 1]
    if prev.tag == f"{{{W}}}p":
        prev_ppr = prev.find(f"{{{W}}}pPr")
        if prev_ppr is not None and prev_ppr.find(f"{{{W}}}sectPr") is not None:
            return False

    break_para = ET.Element(f"{{{W}}}p")
    ppr = ET.SubElement(break_para, f"{{{W}}}pPr")
    sectpr = deepcopy(body_sectpr)
    pg_num = sectpr.find(f"{{{W}}}pgNumType")
    if pg_num is None:
        pg_num = ET.SubElement(sectpr, f"{{{W}}}pgNumType")
    pg_num.set(f"{{{W}}}fmt", "decimal")
    pg_num.set(f"{{{W}}}start", "1")
    ppr.append(sectpr)

    body.insert(target_idx, break_para)
    return True


def set_geometry(docx: Path) -> Counter:
    doc = Document(docx)
    stats: Counter = Counter()

    for idx, section in enumerate(doc.sections, start=1):
        if section.orientation == WD_ORIENT.LANDSCAPE:
            section.page_width, section.page_height = LANDSCAPE_PAGE
            margins = LANDSCAPE_MARGINS
            stats["landscape_sections"] += 1
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = PORTRAIT_PAGE
            margins = PORTRAIT_MARGINS
            stats["portrait_sections"] += 1

        section.top_margin = margins["top"]
        section.bottom_margin = margins["bottom"]
        section.left_margin = margins["left"]
        section.right_margin = margins["right"]
        section.header_distance = HEADER_DISTANCE
        section.footer_distance = FOOTER_DISTANCE
        section.gutter = GUTTER

    doc.settings.odd_and_even_pages_header_footer = False
    doc.save(docx)
    return stats


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def append_page_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_begin = OxmlElement("w:r")
    r_begin.append(fld_begin)
    paragraph._p.append(r_begin)

    r_instr = OxmlElement("w:r")
    r_instr.append(instr_text)
    paragraph._p.append(r_instr)

    r_sep = OxmlElement("w:r")
    r_sep.append(fld_separate)
    paragraph._p.append(r_sep)

    r_end = OxmlElement("w:r")
    r_end.append(fld_end)
    paragraph._p.append(r_end)


def ensure_footer_page_fields(docx: Path) -> Counter:
    doc = Document(docx)
    stats: Counter = Counter()

    for idx, section in enumerate(doc.sections, start=1):
        section.different_first_page_header_footer = False
        footer = section.footer
        footer.is_linked_to_previous = False

        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph_runs(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Keep the cover/title preliminary section footer blank.
        if idx == 1:
            stats["footer_blank_cover_sections"] += 1
            continue

        append_page_field(paragraph)
        stats["footer_page_field_sections"] += 1

    doc.save(docx)
    return stats


def ensure_reference_heading_gap(docx: Path) -> Counter:
    doc = Document(docx)
    stats: Counter = Counter()

    for idx, paragraph in enumerate(doc.paragraphs):
        text = "".join(paragraph.text.split())
        if not (paragraph.style and paragraph.style.name == "Heading 1" and text == "參考文獻"):
            continue

        p_pr = paragraph._p.get_or_add_pPr()
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "240")
        stats["reference_heading_gap_adjusted"] += 1
        break

    doc.save(docx)
    return stats


def ensure_pg_num_type(sect_pr: ET.Element) -> ET.Element:
    pg_num = sect_pr.find(f"{{{W}}}pgNumType")
    if pg_num is None:
        pg_num = ET.SubElement(sect_pr, f"{{{W}}}pgNumType")
    return pg_num


def remove_pg_num_type(sect_pr: ET.Element) -> None:
    pg_num = sect_pr.find(f"{{{W}}}pgNumType")
    if pg_num is not None:
        sect_pr.remove(pg_num)


def resolve_section_indices(root: ET.Element, total_sections: int) -> tuple[int, int]:
    body = root.find(f".//{{{W}}}body")
    if body is None:
        return (2 if total_sections >= 2 else 1, 3 if total_sections >= 3 else total_sections)

    current_section_index = 1
    roman_section_index: int | None = None
    chapter_section_index: int | None = None

    for node in list(body):
        if node.tag == f"{{{W}}}p":
            text = compact_text(paragraph_text(node))
            if text:
                if roman_section_index is None and text in ROMAN_START_MARKERS:
                    roman_section_index = current_section_index
                if chapter_section_index is None and text in CHAPTER_START_MARKERS:
                    chapter_section_index = current_section_index

            ppr = node.find(f"{{{W}}}pPr")
            if ppr is not None and ppr.find(f"{{{W}}}sectPr") is not None:
                current_section_index += 1

    if roman_section_index is None:
        roman_section_index = 2 if total_sections >= 2 else 1

    if chapter_section_index is None:
        chapter_section_index = roman_section_index + 1

    roman_section_index = max(1, min(roman_section_index, total_sections))
    chapter_section_index = max(roman_section_index + 1, min(chapter_section_index, total_sections))
    return roman_section_index, chapter_section_index


def set_page_numbering(docx: Path) -> Counter:
    with zipfile.ZipFile(docx) as zf:
        xml = zf.read("word/document.xml")

    root = ET.fromstring(xml)
    inserted = ensure_main_matter_section_break(root)
    sect_prs = root.findall(".//w:sectPr", NS)
    stats: Counter = Counter()
    roman_section_index, arabic_section_index = resolve_section_indices(root, len(sect_prs))
    stats["resolved_roman_section_index"] = roman_section_index
    stats["resolved_arabic_section_index"] = arabic_section_index
    if inserted:
        stats["inserted_main_matter_section_break"] += 1

    for idx, sect_pr in enumerate(sect_prs, start=1):
        if idx == roman_section_index:
            pg_num = ensure_pg_num_type(sect_pr)
            pg_num.set(f"{{{W}}}fmt", "lowerRoman")
            pg_num.set(f"{{{W}}}start", "1")
            stats["roman_sections"] += 1
        elif roman_section_index < idx < arabic_section_index:
            pg_num = ensure_pg_num_type(sect_pr)
            pg_num.set(f"{{{W}}}fmt", "lowerRoman")
            pg_num.attrib.pop(f"{{{W}}}start", None)
            stats["roman_sections"] += 1
        elif idx == arabic_section_index:
            pg_num = ensure_pg_num_type(sect_pr)
            pg_num.set(f"{{{W}}}fmt", "decimal")
            pg_num.set(f"{{{W}}}start", "1")
            stats["arabic_restart_sections"] += 1
        elif idx > arabic_section_index:
            # Keep decimal numbering continuous after Chapter 1.
            pg_num = ensure_pg_num_type(sect_pr)
            pg_num.set(f"{{{W}}}fmt", "decimal")
            pg_num.attrib.pop(f"{{{W}}}start", None)
        else:
            remove_pg_num_type(sect_pr)

    write_docx_xml(docx, ET.ElementTree(root))
    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize thesis page geometry and page-number section rules.")
    parser.add_argument("--docx", required=True, type=Path)
    args = parser.parse_args()

    geometry_stats = set_geometry(args.docx)
    footer_stats = ensure_footer_page_fields(args.docx)
    reference_gap_stats = ensure_reference_heading_gap(args.docx)
    numbering_stats = set_page_numbering(args.docx)
    stats = geometry_stats + footer_stats + reference_gap_stats + numbering_stats
    for key in sorted(stats):
        print(f"{key}={stats[key]}")


if __name__ == "__main__":
    main()
