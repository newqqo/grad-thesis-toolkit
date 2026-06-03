from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from thesis_render_v2 import (
    DEFAULT_SHADOW,
    chapter_heading_text,
    chapter_title_text,
    ensure_v2_styles,
    first_body_heading,
    is_body_chapter_heading,
    load_p_style_overrides,
    load_shadow_chapters,
    load_style_profile,
    normalize_style_alias,
    resolve_chapter_style_rule,
    visible_text,
    DEFAULT_STYLE_PROFILE,
)


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "deliverables" / "docx" / "thesis_render_v2_latest.docx"
DEFAULT_P_STYLE_OVERRIDES = ROOT / "consistency" / "rules" / "p_style_overrides.json"
FRONT_MATTER_HEADING_TEXTS = {"中文摘要", "Abstract", "誌謝", "謝誌", "目錄", "表目錄", "圖目錄"}
APPENDIX_CHAPTER_ID = "appendices"
APPENDIX_MAIN_HEADING_RE = re.compile(r"^附錄\s+[A-ZＡ-Ｚ](?:\s|$)")
APPENDIX_H3_RE = re.compile(r"^(?:[A-Z]\.\d+(?![A-Za-z0-9])|[一二三四五六七八九十]+、)")
APPENDIX_H4_RE = re.compile(r"^(?:[A-Z]\.\d+[a-z](?![A-Za-z0-9])|（[一二三四五六七八九十]+）)")


@dataclass
class ExpectedParagraphStyle:
    p_index: int
    text: str
    style_name: str


def normalize_for_match(text: str) -> str:
    normalized = (text or "").strip()
    normalized = re.sub(r"^\d+[\.|\)|、]\s*", "", normalized)
    normalized = re.sub(r"^\(\d+\)\s*", "", normalized)
    return re.sub(r"\s+", "", normalized).strip()


def is_appendix_heading_text(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return bool(
        APPENDIX_MAIN_HEADING_RE.match(stripped)
        or APPENDIX_H3_RE.match(stripped)
        or APPENDIX_H4_RE.match(stripped)
    )


def build_expected_body_styles(style_profile: dict) -> list[ExpectedParagraphStyle]:
    fallback_style_name = str(style_profile.get("fallback_style", "Body Normal")).strip() or "Body Normal"
    p_style_overrides = load_p_style_overrides(DEFAULT_P_STYLE_OVERRIDES)
    chapters = load_shadow_chapters(DEFAULT_SHADOW)

    expected: list[ExpectedParagraphStyle] = []
    for chapter in chapters:
        if not (chapter.chapter_id.startswith("ch") or chapter.chapter_id == APPENDIX_CHAPTER_ID):
            continue

        skip_first = first_body_heading(chapter)
        paragraphs = chapter.paragraphs
        if skip_first and paragraphs and visible_text(paragraphs[0].text).strip() in {
            skip_first,
            chapter.label,
            chapter_heading_text(chapter),
            chapter_title_text(chapter.label),
        }:
            paragraphs = paragraphs[1:]

        for paragraph_record in paragraphs:
            stripped = visible_text(paragraph_record.text).strip()
            if not stripped:
                continue
            # Keep appendix heading-style paragraphs out of this repair map to
            # avoid touching TOC-sensitive presentation logic.
            if chapter.chapter_id == APPENDIX_CHAPTER_ID and is_appendix_heading_text(stripped):
                continue
            override_style = normalize_style_alias(p_style_overrides.get(paragraph_record.index, ""))
            if override_style:
                style_name = override_style
            else:
                style_name = resolve_chapter_style_rule(chapter.chapter_id, stripped, style_profile) or fallback_style_name
            expected.append(ExpectedParagraphStyle(paragraph_record.index, stripped, style_name))
    return expected


def find_body_start_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_body_chapter_heading(paragraph.text.strip()):
            return idx
    raise ValueError("Could not find imported v2 chapter start in merged DOCX.")


def resolve_style_for_paragraph(
    text: str,
    expected: list[ExpectedParagraphStyle],
    cursor: int,
    max_lookahead: int = 8,
) -> tuple[str | None, int]:
    normalized_text = normalize_for_match(text)
    if cursor >= len(expected):
        return None, cursor

    if normalize_for_match(expected[cursor].text) == normalized_text:
        return expected[cursor].style_name, cursor + 1

    # Allow short look-ahead to recover from occasional drift after field updates.
    for lookahead in range(cursor + 1, min(cursor + max_lookahead, len(expected))):
        if normalize_for_match(expected[lookahead].text) == normalized_text:
            return expected[lookahead].style_name, lookahead + 1

    return None, cursor


def clear_paragraph_numbering(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


def ensure_front_matter_heading_style(doc: Document) -> str:
    style_name = "Front Matter Heading"
    if style_name in {style.name for style in doc.styles}:
        style = doc.styles[style_name]
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        heading1 = doc.styles["Heading 1"]
        style.base_style = doc.styles["Normal"]
        style.font.name = heading1.font.name
        style.font.size = heading1.font.size
        style.font.bold = heading1.font.bold
        style.paragraph_format.alignment = heading1.paragraph_format.alignment
        style.paragraph_format.space_before = heading1.paragraph_format.space_before
        style.paragraph_format.space_after = heading1.paragraph_format.space_after

    ppr = style.element.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)

    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), "0")
    return style_name


def repair(docx: Path) -> None:
    doc = Document(docx)
    style_profile = load_style_profile(DEFAULT_STYLE_PROFILE)
    ensure_v2_styles(doc, style_profile=style_profile, preserve_front_matter=True)
    expected = build_expected_body_styles(style_profile)

    start_idx = find_body_start_index(doc)
    front_heading_style = ensure_front_matter_heading_style(doc)

    for paragraph in doc.paragraphs[:start_idx]:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name == "Heading 1" and text in FRONT_MATTER_HEADING_TEXTS:
            paragraph.style = doc.styles[front_heading_style]
            clear_paragraph_numbering(paragraph)

    cursor = 0

    for rel_index, paragraph in enumerate(doc.paragraphs[start_idx:], start=0):
        text = paragraph.text.strip()
        if not text:
            continue

        if rel_index == 0 or is_body_chapter_heading(text) or text == "參考文獻":
            paragraph.style = doc.styles["Heading 1"]
            clear_paragraph_numbering(paragraph)
            continue

        if is_appendix_heading_text(text):
            # Preserve current appendix heading presentation to avoid TOC/page-flow side effects.
            clear_paragraph_numbering(paragraph)
            continue

        style_name, cursor = resolve_style_for_paragraph(text, expected, cursor)
        if style_name:
            paragraph.style = doc.styles[style_name]
            clear_paragraph_numbering(paragraph)

    doc.save(docx)


def main() -> None:
    parser = argparse.ArgumentParser(description="Repair imported v2 paragraph styles after Word field update.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()
    repair(args.docx.resolve())


if __name__ == "__main__":
    main()
