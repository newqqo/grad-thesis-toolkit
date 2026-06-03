from __future__ import annotations

import argparse
import ast
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SHADOW = ROOT / "source" / "shadow"
DEFAULT_LAYOUT = ROOT / "source" / "layout" / "thesis_v2_layout.md"
DEFAULT_TEMPLATE = ROOT / "assets" / "templates" / "docx" / "thesis_template_thin.docx"
DEFAULT_OUTPUT = ROOT / "deliverables" / "docx" / "thesis_render_v2_latest.docx"
DEFAULT_STYLE_PROFILE = ROOT / "consistency" / "rules" / "v2_style_profiles.json"
TABLES_DIR = ROOT / "assets" / "tables"

CAPTION_RE = re.compile(r"^(附表|附圖|表|圖)\s*[A-Z]?\d+")
APPENDIX_MAIN_RE = re.compile(r"^附錄\s+[A-ZＡ-Ｚ](?:\s|$)")
CHAPTER_HEADING_RE = re.compile(r"^第[一二三四五六七八九十]+章\s+")
TOP_ITEM_RE = re.compile(r"^\d+\.\s+")
SUB_ITEM_RE = re.compile(r"^\([0-9]+\)\s+")
INTERVIEW_Q_RE = re.compile(r"^Q\d+(?:-\d+[a-zA-Z]?)?\b")
INTERVIEW_Q_LABEL_RE = re.compile(r"^(Q\d+)(\s+.*)$")
STYLE_HINT_RE = re.compile(r"^<!--\s*(h[1-5]|appendix|caption|note|l4|l4lead|l4body|l5|l5body)\s*-->\s*$")
SHORT_HEADING_MAX = 28
PAGE_WIDTH = Inches(8.27)
PAGE_HEIGHT = Inches(11.69)
LEGACY_STYLE_SOURCE = ROOT / "deliverables" / "docx" / "thesis_render_latest.docx"
NOTE_LINE_INDENT = Pt(54)
SOURCE_LINE_INDENT = Pt(54)

DEFAULT_STYLE_PROFILE_DATA = {
    "fallback_style": "Body Normal",
    "heading_numbering": {
        "enabled": False,
        "levels": [2, 3],
        "separator": ".",
        "native_layout": {
            "level_1": {"left_indent_pt": 0, "hanging_indent_pt": 0},
            "level_2": {"left_indent_pt": 0, "hanging_indent_pt": 0},
            "level_3": {"left_indent_pt": 0, "hanging_indent_pt": 0},
        },
    },
    "font_policy": {
        "ascii": "Times New Roman",
        "hAnsi": "Times New Roman",
        "eastAsia": "新細明體",
        "cs": "Times New Roman",
    },
    "table_layout": {
        "portrait_textbox_twips": round((21.0 - 3.0 - 3.0) / 2.54 * 1440),
        "landscape_textbox_twips": round((29.7 - 3.0 - 3.0) / 2.54 * 1440),
        "border": {"val": "single", "sz": "4", "space": "0", "color": "auto"},
    },
    # Styles are JSON-driven from consistency/rules/v2_style_profiles.json.
    "styles": {},
}


@dataclass
class ShadowChapter:
    chapter_id: str
    label: str
    paragraphs: list["ShadowParagraph"]


@dataclass
class ShadowParagraph:
    index: int
    text: str


@dataclass
class TableDef:
    table_id: str
    title_match: str
    title: str
    rows: list[list[str]]
    col_widths: list[int]
    font_size: int
    page_layout: str = "portrait"


@dataclass
class LayoutRule:
    kind: str
    match: re.Pattern[str]
    path: Path | None = None
    width_in: float | None = None
    script: Path | None = None
    style_name: str | None = None


def parse_scalar(value: str):
    value = value.strip()
    if not value:
        return ""
    try:
        return ast.literal_eval(value)
    except Exception:
        return value.strip('"')


def split_frontmatter(text: str) -> tuple[dict[str, object], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---", 4)
    if end == -1:
        return {}, text
    fm_text = text[4:end].strip()
    body = text[end + 4 :].lstrip("\r\n")
    data: dict[str, object] = {}
    for raw_line in fm_text.splitlines():
        line = raw_line.strip()
        if not line or ":" not in line:
            continue
        key, value = line.split(":", 1)
        data[key.strip()] = parse_scalar(value)
    return data, body


def parse_shadow_file(path: Path) -> list[ShadowParagraph]:
    _, body = split_frontmatter(path.read_text(encoding="utf-8"))
    paragraphs: list[ShadowParagraph] = []
    current: list[str] = []
    current_index: int | None = None
    in_comment = False
    in_block = False
    for raw_line in body.splitlines():
        stripped = raw_line.strip()
        if stripped.startswith("<!--") and STYLE_HINT_RE.match(stripped):
            if in_block:
                current.append(raw_line.rstrip())
            continue
        if stripped.startswith("<!--"):
            in_comment = True
        if in_comment:
            if "-->" in stripped:
                in_comment = False
            continue
        if stripped.startswith("# "):
            continue
        marker = re.match(r"^## p(\d+)$", stripped)
        if marker:
            if in_block:
                text = "\n".join(current).strip()
                if text and current_index is not None:
                    paragraphs.append(ShadowParagraph(index=current_index, text=text))
            current = []
            in_block = True
            current_index = int(marker.group(1))
            continue
        if in_block:
            current.append(raw_line.rstrip())
    if in_block:
        text = "\n".join(current).strip()
        if text and current_index is not None:
            paragraphs.append(ShadowParagraph(index=current_index, text=text))
    return paragraphs


def load_shadow_chapters(shadow_dir: Path) -> list[ShadowChapter]:
    manifest_path = shadow_dir / "shadow_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    chapters: list[ShadowChapter] = []
    for chapter in manifest["chapters"]:
        chapter_id = chapter["id"]
        path = shadow_dir / f"{chapter_id}.md"
        chapters.append(
            ShadowChapter(
                chapter_id=chapter_id,
                label=chapter["label"],
                paragraphs=parse_shadow_file(path),
            )
        )
    return chapters


def normalize_style_alias(style_name: str) -> str:
    key = style_name.strip().lower().replace("_", "").replace(" ", "")
    aliases = {
        "h1": "Heading 1",
        "head1": "Heading 1",
        "heading1": "Heading 1",
        "h2": "Heading 2",
        "head2": "Heading 2",
        "heading2": "Heading 2",
        "h3": "Heading 3",
        "head3": "Heading 3",
        "heading3": "Heading 3",
        "h4": "Heading 4",
        "head4": "Heading 4",
        "heading4": "Heading 4",
        "h5": "Heading 5",
        "head5": "Heading 5",
        "heading5": "Heading 5",
    }
    return aliases.get(key, style_name.strip())


def load_p_style_overrides(path: Path) -> dict[int, str]:
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    entries = data.get("styles", {}) if isinstance(data, dict) else {}
    if not isinstance(entries, dict):
        return {}
    overrides: dict[int, str] = {}
    for p_text, raw_style in entries.items():
        try:
            p_index = int(str(p_text))
        except ValueError:
            continue
        style_name = normalize_style_alias(str(raw_style))
        if not style_name:
            continue
        overrides[p_index] = style_name
    return overrides


def load_p_section_break_before_overrides(path: Path) -> dict[int, str]:
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        return {}

    raw = data.get("section_break_before", {})
    if not isinstance(raw, dict):
        return {}

    overrides: dict[int, str] = {}
    for p_text, value in raw.items():
        try:
            p_index = int(str(p_text))
        except ValueError:
            continue
        orientation = str(value).strip().lower()
        if orientation not in {"portrait", "landscape"}:
            orientation = "portrait"
        overrides[p_index] = orientation
    return overrides


def merge_dict(base: dict, override: dict) -> dict:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = merge_dict(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_style_profile(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"Style profile not found: {path}")
    loaded = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(loaded, dict):
        raise ValueError(f"Style profile must be a JSON object: {path}")
    return merge_dict(DEFAULT_STYLE_PROFILE_DATA, loaded)


def resolve_font_policy(style_profile: dict) -> dict[str, str]:
    defaults = DEFAULT_STYLE_PROFILE_DATA.get("font_policy", {})
    raw = style_profile.get("font_policy", {}) if isinstance(style_profile, dict) else {}
    policy = dict(defaults)
    if isinstance(raw, dict):
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            value = raw.get(key)
            if isinstance(value, str) and value.strip():
                policy[key] = value.strip()
    return policy


def set_rfonts(rpr, font_map: dict[str, str]) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for theme_key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        attr = qn(f"w:{theme_key}")
        if attr in rfonts.attrib:
            del rfonts.attrib[attr]
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = font_map.get(key)
        if value:
            rfonts.set(qn(f"w:{key}"), value)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(cell) <= {"-", " "} for cell in rows[1]):
        rows.pop(1)
    return rows


def load_table_defs(tables_dir: Path) -> list[TableDef]:
    table_defs: list[TableDef] = []
    for path in sorted(tables_dir.glob("table_*.md")):
        frontmatter, body = split_frontmatter(path.read_text(encoding="utf-8"))
        title = ""
        table_lines: list[str] = []
        for line in body.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
            elif line.strip().startswith("|"):
                table_lines.append(line)
        table_defs.append(
            TableDef(
                table_id=str(frontmatter.get("table_id", path.stem)),
                title_match=str(frontmatter.get("title_match", title)),
                title=title,
                rows=parse_markdown_table(table_lines),
                col_widths=list(frontmatter.get("col_widths", [])),
                font_size=int(frontmatter.get("font_size", 10)),
                page_layout=str(frontmatter.get("page_layout", "portrait")),
            )
        )
    return table_defs


def parse_kv_pairs(line: str) -> dict[str, str]:
    matches = re.findall(r'([a-zA-Z_]+)="([^"]*)"', line)
    return {key: value for key, value in matches}


def load_layout_rules(path: Path) -> list[LayoutRule]:
    if not path.exists():
        return []
    rules: list[LayoutRule] = []
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line.startswith(":::"):
            continue
        payload = line[3:].strip()
        kind = payload.split(" ", 1)[0]
        kv = parse_kv_pairs(payload)
        match_text = kv.get("match")
        if not match_text:
            continue
        rule = LayoutRule(
            kind=kind,
            match=re.compile(match_text),
            path=(ROOT / kv["path"]).resolve() if "path" in kv else None,
            width_in=float(kv["width_in"]) if "width_in" in kv else None,
            script=(ROOT / kv["script"]).resolve() if "script" in kv else None,
            style_name=kv.get("style"),
        )
        rules.append(rule)
    return rules


def strip_style_hints(text: str) -> tuple[str | None, str]:
    lines = text.splitlines()
    hint: str | None = None
    while lines:
        match = STYLE_HINT_RE.match(lines[0].strip())
        if not match:
            break
        hint = match.group(1).lower()
        lines = lines[1:]
    return hint, "\n".join(lines).strip()


def visible_text(text: str) -> str:
    return strip_style_hints(text)[1]


def style_hint_to_style_name(hint: str | None) -> str | None:
    mapping = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
        "h4": "Heading 4",
        "h5": "Heading 5",
        "appendix": "Appendix Heading",
        "caption": "Caption",
        "note": "圖表註",
        "l4": "Thesis Level 4",
        "l4lead": "Thesis Level 4 Lead",
        "l4body": "Thesis Level 4 Body",
        "l5": "Thesis Level 5",
        "l5body": "Thesis Level 5 Body",
    }
    return mapping.get(hint or "")


def chapter_title_text(label: str) -> str:
    parts = label.split(" ", 1)
    if label.startswith("第") and len(parts) == 2:
        return parts[1].strip()
    return label.strip()


def body_chapter_titles() -> set[str]:
    manifest_path = DEFAULT_SHADOW / "shadow_manifest.json"
    if not manifest_path.exists():
        return set()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    return {chapter_title_text(chapter["label"]) for chapter in manifest["chapters"] if chapter["id"].startswith("ch")}


def is_body_chapter_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    normalized = re.sub(r"^[一二三四五六七八九十]+、\s*", "", stripped)
    return bool(stripped) and (
        CHAPTER_HEADING_RE.match(stripped)
        or stripped in body_chapter_titles()
        or normalized in body_chapter_titles()
    )


def load_legacy_style_map(path: Path = LEGACY_STYLE_SOURCE) -> dict[str, str]:
    if not path.exists():
        return {}
    doc = Document(path)
    allowed = {
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Caption",
        "圖表註",
        "Appendix Heading",
        "Thesis Level 4",
        "Thesis Level 4 Lead",
        "Thesis Level 4 Body",
        "Thesis Level 5",
        "Thesis Level 5 Body",
    }
    style_map: dict[str, str] = {}
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if text and style_name in allowed and text not in style_map:
            style_map[text] = style_name
    return style_map


def remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def ensure_style(doc: Document, style_name: str, base_style: str = "Normal"):
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style]
        return style


def resolve_body_style_settings(styles_cfg: dict, fallback_style_name: str) -> dict:
    # Backward-compatible lookup: use explicitly selected fallback style first,
    # then profile's Normal block, finally legacy No Spacing block.
    if fallback_style_name in styles_cfg and isinstance(styles_cfg[fallback_style_name], dict):
        return styles_cfg[fallback_style_name]
    if "Normal" in styles_cfg and isinstance(styles_cfg["Normal"], dict):
        return styles_cfg["Normal"]
    if "No Spacing" in styles_cfg and isinstance(styles_cfg["No Spacing"], dict):
        return styles_cfg["No Spacing"]
    return {}


def apply_style_profile(style, settings: dict, font_policy: dict[str, str]) -> None:
    style_font_map = dict(font_policy)
    style_font_override = settings.get("rfonts", {})
    if isinstance(style_font_override, dict):
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            value = style_font_override.get(key)
            if isinstance(value, str) and value.strip():
                style_font_map[key] = value.strip()

    style_rpr = style.element.get_or_add_rPr()
    set_rfonts(style_rpr, style_font_map)

    style.font.name = style_font_map.get("ascii", style.font.name)
    if "font_name" in settings:
        style.font.name = settings["font_name"]
    if "font_size_pt" in settings:
        style.font.size = Pt(float(settings["font_size_pt"]))
    if "bold" in settings:
        style.font.bold = bool(settings["bold"])

    pf = style.paragraph_format
    if "space_before_pt" in settings:
        pf.space_before = Pt(float(settings["space_before_pt"]))
    if "space_after_pt" in settings:
        pf.space_after = Pt(float(settings["space_after_pt"]))
    if "left_indent_pt" in settings:
        pf.left_indent = Pt(float(settings["left_indent_pt"]))
    if "first_line_indent_pt" in settings:
        pf.first_line_indent = Pt(float(settings["first_line_indent_pt"]))

    # Word-like special indent control via JSON:
    # - special_indent: none | first_line | hanging
    # - special_indent_by_pt: numeric point value for the dropdown's "By" field
    indent_mode_raw = settings.get("special_indent")
    if indent_mode_raw is not None:
        indent_mode = str(indent_mode_raw).strip().lower()
        indent_by = float(settings.get("special_indent_by_pt", 0))
        indent_by = abs(indent_by)
        if indent_mode in {"none", "無"}:
            pf.first_line_indent = Pt(0)
        elif indent_mode in {"first_line", "firstline", "首行"}:
            pf.first_line_indent = Pt(indent_by)
        elif indent_mode in {"hanging", "凸排"}:
            # Hanging indent is represented by negative first-line indent.
            pf.first_line_indent = Pt(-indent_by)
            if "left_indent_pt" not in settings:
                pf.left_indent = Pt(indent_by)
    if "line_spacing_pt" in settings:
        pf.line_spacing = Pt(float(settings["line_spacing_pt"]))

    rule_map = {
        "single": WD_LINE_SPACING.SINGLE,
        "one_point_five": WD_LINE_SPACING.ONE_POINT_FIVE,
        "double": WD_LINE_SPACING.DOUBLE,
        "at_least": WD_LINE_SPACING.AT_LEAST,
        "exact": WD_LINE_SPACING.EXACTLY,
        "multiple": WD_LINE_SPACING.MULTIPLE,
    }
    if "line_spacing_rule" in settings:
        rule_key = str(settings["line_spacing_rule"]).strip().lower()
        if rule_key in rule_map:
            pf.line_spacing_rule = rule_map[rule_key]

    if "line_spacing" in settings and "line_spacing_pt" not in settings:
        pf.line_spacing = float(settings["line_spacing"])
    if "keep_with_next" in settings:
        pf.keep_with_next = bool(settings["keep_with_next"])
    if "page_break_before" in settings:
        pf.page_break_before = bool(settings["page_break_before"])
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distributed": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    if "alignment" in settings:
        key = str(settings["alignment"]).strip().lower()
        if key in alignment_map:
            pf.alignment = alignment_map[key]


def ensure_v2_styles(doc: Document, style_profile: dict | None = None, preserve_front_matter: bool = False) -> None:
    if style_profile is None:
        style_profile = load_style_profile(DEFAULT_STYLE_PROFILE)
    styles_cfg = style_profile.get("styles", {})
    font_policy = resolve_font_policy(style_profile)
    fallback_style_name = str(style_profile.get("fallback_style", "Body Normal")).strip() or "Body Normal"

    if not preserve_front_matter:
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_rpr = style.element.get_or_add_rPr()
                set_rfonts(style_rpr, font_policy)

    if not preserve_front_matter or fallback_style_name.strip().lower() != "normal":
        body_style = ensure_style(doc, fallback_style_name)
        apply_style_profile(body_style, resolve_body_style_settings(styles_cfg, fallback_style_name), font_policy)

    appendix = ensure_style(doc, "Appendix Heading")
    appendix_cfg = styles_cfg.get("Appendix Heading", {})
    apply_style_profile(appendix, appendix_cfg, font_policy)
    # outlineLvl=0 so appendix headings appear in TOC at chapter level
    pPr = appendix.element.get_or_add_pPr()
    existing_ol = pPr.find(qn("w:outlineLvl"))
    appendix_outline = int(appendix_cfg.get("outline_level", 0))
    if existing_ol is None:
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), str(appendix_outline))
        pPr.append(ol)
    else:
        existing_ol.set(qn("w:val"), str(appendix_outline))

    if not preserve_front_matter:
        # TOC styles: non-bold, proper indentation
        toc_cfg = styles_cfg.get("TOC", {})
        toc_indents = [Pt(float(item)) for item in toc_cfg.get("level_left_indent_pt", [0, 12, 24, 36])]
        for level in range(1, 5):
            toc_s = ensure_style(doc, f"TOC {level}")
            apply_style_profile(toc_s, toc_cfg, font_policy)
            toc_s.paragraph_format.left_indent = toc_indents[level - 1]
            toc_s.paragraph_format.first_line_indent = Pt(float(toc_cfg.get("first_line_indent_pt", 0)))

        # Some templates use lowercase/localized TOC style names; normalize all variants.
        toc_name_set = {str(name).strip().lower() for name in toc_cfg.get("variant_names", [])}
        for style in doc.styles:
            style_name = (style.name or "").strip().lower()
            if style_name not in toc_name_set:
                continue
            level = 1
            level_match = re.search(r"([1-4])$", style_name)
            if level_match:
                level = int(level_match.group(1))
            apply_style_profile(style, toc_cfg, font_policy)
            style.paragraph_format.left_indent = toc_indents[level - 1]
            style.paragraph_format.first_line_indent = Pt(float(toc_cfg.get("first_line_indent_pt", 0)))

    note = ensure_style(doc, "圖表註")
    apply_style_profile(note, styles_cfg.get("圖表註", {}), font_policy)

    for heading_style in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"]:
        if heading_style in styles_cfg:
            style = ensure_style(doc, heading_style)
            apply_style_profile(style, styles_cfg.get(heading_style, {}), font_policy)

    for sub in ("圖", "表"):
        csub = ensure_style(doc, f"Caption {sub}")
        try:
            csub.base_style = doc.styles["Caption"]
        except KeyError:
            pass

    for style_name in [
        "Thesis Level 4",
        "Thesis Level 4 Lead",
        "Thesis Level 4 Body",
        "Thesis Level 5",
        "Thesis Level 5 Body",
        "Reference Entry",
    ]:
        style = ensure_style(doc, style_name)
        style.base_style = doc.styles["Normal"]
        apply_style_profile(style, styles_cfg.get(style_name, {}), font_policy)

    ensure_native_heading_numbering(doc, font_policy, style_profile)
    ensure_native_thesis_item_numbering(doc, font_policy, style_profile)


def resolve_heading_native_layout(style_profile: dict) -> dict[int, dict[str, float]]:
    heading_cfg = style_profile.get("heading_numbering", {})
    native_layout = heading_cfg.get("native_layout", {})
    if not isinstance(native_layout, dict):
        native_layout = {}

    result: dict[int, dict[str, float]] = {}
    for ilvl in range(3):
        idx = ilvl + 1
        candidates = [
            native_layout.get(f"level_{idx}"),
            native_layout.get(f"heading_{idx}"),
            native_layout.get(f"h{idx}"),
        ]
        selected = next((item for item in candidates if isinstance(item, dict)), {})
        cfg = {
            "left_indent_pt": float(selected.get("left_indent_pt", 0)),
            "hanging_indent_pt": float(selected.get("hanging_indent_pt", 0)),
        }
        if "first_line_indent_pt" in selected:
            cfg["first_line_indent_pt"] = float(selected.get("first_line_indent_pt", 0))
        result[ilvl] = cfg
    return result


def apply_level_indent(lvl, layout_cfg: dict[str, float]) -> None:
    ppr = lvl.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        lvl.append(ppr)

    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)

    left_twips = int(round(float(layout_cfg.get("left_indent_pt", 0)) * 20))
    hanging_twips = int(round(float(layout_cfg.get("hanging_indent_pt", 0)) * 20))
    ind.set(qn("w:left"), str(max(0, left_twips)))

    if "first_line_indent_pt" in layout_cfg:
        first_line_twips = int(round(float(layout_cfg.get("first_line_indent_pt", 0)) * 20))
        ind.set(qn("w:firstLine"), str(first_line_twips))
        if qn("w:hanging") in ind.attrib:
            del ind.attrib[qn("w:hanging")]
    else:
        ind.set(qn("w:hanging"), str(max(0, hanging_twips)))
        if qn("w:firstLine") in ind.attrib:
            del ind.attrib[qn("w:firstLine")]


def ensure_native_heading_numbering(doc: Document, font_policy: dict[str, str], style_profile: dict) -> None:
    numbering = doc.part.numbering_part.element
    abs_id = "900"
    num_id = "900"

    abstract = None
    for item in numbering.findall(qn("w:abstractNum")):
        if item.get(qn("w:abstractNumId")) == abs_id:
            abstract = item
            break

    if abstract is None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abs_id)
        numbering.append(abstract)

    # Rebuild levels each render to keep numbering deterministic.
    for child in list(abstract):
        if child.tag in {qn("w:lvl"), qn("w:multiLevelType")}:
            abstract.remove(child)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)

    level_specs = [
        (0, "Heading1", "chineseCounting", "%1、", False),
        (1, "Heading2", "decimal", "%1.%2", True),
        (2, "Heading3", "decimal", "%1.%2.%3", True),
    ]
    level_layouts = resolve_heading_native_layout(style_profile)
    for ilvl, pstyle, numfmt, lvl_text, legal in level_specs:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), numfmt)
        lvl.append(num_fmt)

        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), pstyle)
        lvl.append(p_style)

        lvl_text_el = OxmlElement("w:lvlText")
        lvl_text_el.set(qn("w:val"), lvl_text)
        lvl.append(lvl_text_el)

        if legal:
            lvl.append(OxmlElement("w:isLgl"))

        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "nothing" if ilvl == 0 else "space")
        lvl.append(suff)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        apply_level_indent(lvl, level_layouts.get(ilvl, {"left_indent_pt": 0, "hanging_indent_pt": 0}))

        lvl_rpr = OxmlElement("w:rPr")
        set_rfonts(lvl_rpr, font_policy)
        lvl.append(lvl_rpr)

        abstract.append(lvl)

    num = None
    for item in numbering.findall(qn("w:num")):
        if item.get(qn("w:numId")) == num_id:
            num = item
            break

    if num is None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        numbering.append(num)

    # Force numId -> abstractNumId mapping every render so stale template bindings
    # cannot silently redirect Heading numbering to a different list chain.
    abs_ref = num.find(qn("w:abstractNumId"))
    if abs_ref is None:
        abs_ref = OxmlElement("w:abstractNumId")
        num.insert(0, abs_ref)
    abs_ref.set(qn("w:val"), abs_id)

    style_level_map = [
        ("Heading 1", 0),
        ("Heading 2", 1),
        ("Heading 3", 2),
    ]
    for style_name, ilvl in style_level_map:
        style = ensure_style(doc, style_name)
        ppr = style.element.get_or_add_pPr()
        numpr = ppr.find(qn("w:numPr"))
        if numpr is None:
            numpr = OxmlElement("w:numPr")
            ppr.append(numpr)

        ilvl_el = numpr.find(qn("w:ilvl"))
        if ilvl_el is None:
            ilvl_el = OxmlElement("w:ilvl")
            numpr.append(ilvl_el)
        ilvl_el.set(qn("w:val"), str(ilvl))

        numid_el = numpr.find(qn("w:numId"))
        if numid_el is None:
            numid_el = OxmlElement("w:numId")
            numpr.append(numid_el)
        numid_el.set(qn("w:val"), num_id)


def strip_level_item_prefix(text: str, style_name: str) -> str:
    stripped = text.strip()
    if style_name == "Thesis Level 4":
        stripped = re.sub(r"^\d+[\.|\)|、]\s*", "", stripped)
    elif style_name == "Thesis Level 5":
        stripped = re.sub(r"^\(\d+\)\s*", "", stripped)
    return stripped


def is_thesis_item_numbering_enabled(style_profile: dict) -> bool:
    cfg = style_profile.get("thesis_item_numbering", {})
    if not isinstance(cfg, dict):
        return False
    return cfg.get("enabled", True) is not False


def resolve_section_policy(style_profile: dict) -> dict:
    defaults = {
        "break_before_chapter_heading": True,
        "break_before_heading_1": True,
        "break_before_appendix_heading": True,
        "allow_table_under_appendix_heading": True,
        "break_before_all_captions": False,
        "break_after_source_note": False,
        "skip_caption_break_if_previous_was_source_break": True,
        "ensure_landscape_sections_for_wide_tables": True,
        "ensure_landscape_sections_for_layout_images": True,
        "dedupe_adjacent_section_breaks": True,
    }
    cfg = style_profile.get("section_policy", {}) if isinstance(style_profile, dict) else {}
    if not isinstance(cfg, dict):
        cfg = {}
    resolved = dict(defaults)
    for key in defaults:
        if key in cfg:
            resolved[key] = bool(cfg[key])
    return resolved


def ensure_native_thesis_item_numbering(doc: Document, font_policy: dict[str, str], style_profile: dict) -> None:
    cfg = style_profile.get("thesis_item_numbering", {})
    if not isinstance(cfg, dict):
        cfg = {}
    if cfg.get("enabled", True) is False:
        return

    numbering = doc.part.numbering_part.element
    abs_id = str(cfg.get("abstract_num_id", 901))
    num_id = str(cfg.get("num_id", 901))

    abstract = None
    for item in numbering.findall(qn("w:abstractNum")):
        if item.get(qn("w:abstractNumId")) == abs_id:
            abstract = item
            break

    if abstract is None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abs_id)
        numbering.append(abstract)

    for child in list(abstract):
        if child.tag in {qn("w:lvl"), qn("w:multiLevelType")}:
            abstract.remove(child)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract.append(multi_level)

    lvl0_text = str(cfg.get("level_4_text", "%1."))
    lvl1_text = str(cfg.get("level_5_text", "(%2)"))

    level_specs = [
        (0, "ThesisLevel4", lvl0_text, "space", None),
        (1, "ThesisLevel5", lvl1_text, "space", 1),
    ]
    for ilvl, pstyle, lvl_text, suff_value, restart_level in level_specs:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)

        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), pstyle)
        lvl.append(p_style)

        lvl_text_el = OxmlElement("w:lvlText")
        lvl_text_el.set(qn("w:val"), lvl_text)
        lvl.append(lvl_text_el)

        if restart_level is not None:
            lvl_restart = OxmlElement("w:lvlRestart")
            lvl_restart.set(qn("w:val"), str(restart_level))
            lvl.append(lvl_restart)

        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), suff_value)
        lvl.append(suff)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        lvl_rpr = OxmlElement("w:rPr")
        set_rfonts(lvl_rpr, font_policy)
        lvl.append(lvl_rpr)

        abstract.append(lvl)

    num = None
    for item in numbering.findall(qn("w:num")):
        if item.get(qn("w:numId")) == num_id:
            num = item
            break

    if num is None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        numbering.append(num)

    # Keep thesis item numbering deterministic across templates/imported docs.
    abs_ref = num.find(qn("w:abstractNumId"))
    if abs_ref is None:
        abs_ref = OxmlElement("w:abstractNumId")
        num.insert(0, abs_ref)
    abs_ref.set(qn("w:val"), abs_id)

    style_level_map = [
        ("Thesis Level 4", 0),
        ("Thesis Level 5", 1),
    ]
    for style_name, ilvl in style_level_map:
        style = ensure_style(doc, style_name)
        ppr = style.element.get_or_add_pPr()
        numpr = ppr.find(qn("w:numPr"))
        if numpr is None:
            numpr = OxmlElement("w:numPr")
            ppr.append(numpr)

        ilvl_el = numpr.find(qn("w:ilvl"))
        if ilvl_el is None:
            ilvl_el = OxmlElement("w:ilvl")
            numpr.append(ilvl_el)
        ilvl_el.set(qn("w:val"), str(ilvl))

        numid_el = numpr.find(qn("w:numId"))
        if numid_el is None:
            numid_el = OxmlElement("w:numId")
            numpr.append(numid_el)
        numid_el.set(qn("w:val"), num_id)


def _apply_section_geometry(section, orientation: str) -> None:
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = PAGE_HEIGHT
        section.page_height = PAGE_WIDTH
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def paragraph_has_section_break(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def add_section(doc: Document, orientation: str, dedupe_adjacent: bool = True) -> None:
    # Guard against back-to-back section breaks on empty paragraphs.
    if dedupe_adjacent and doc.paragraphs:
        last_para = doc.paragraphs[-1]
        if not last_para.text.strip() and paragraph_has_section_break(last_para):
            _apply_section_geometry(doc.sections[-1], orientation)
            return

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _apply_section_geometry(section, orientation)


def looks_like_short_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if CAPTION_RE.match(stripped) or stripped.startswith("註：") or stripped.startswith("資料來源："):
        return False
    if APPENDIX_MAIN_RE.match(stripped):
        return False
    if stripped.endswith(("。", "？", "！", "；", "：")):
        return False
    return len(stripped) <= SHORT_HEADING_MAX


def apply_page_break_rules(paragraph, text: str, rules: list[LayoutRule]) -> None:
    for rule in rules:
        if rule.kind != "page-break-before":
            continue
        if rule.match.search(text):
            paragraph.paragraph_format.page_break_before = True
            break


def matching_image_rule(text: str, rules: list[LayoutRule]) -> LayoutRule | None:
    for rule in rules:
        if rule.kind == "image" and rule.match.search(text):
            return rule
    return None


def twips_to_inches(twips: int) -> float:
    return twips / 1440.0


def resolve_image_max_width_inches(doc: Document, style_profile: dict) -> float:
    section = doc.sections[-1]
    table_layout = style_profile.get("table_layout", {}) if isinstance(style_profile, dict) else {}
    portrait_twips = int(table_layout.get("portrait_textbox_twips", round((21.0 - 3.0 - 3.0) / 2.54 * 1440)))
    landscape_twips = int(table_layout.get("landscape_textbox_twips", round((29.7 - 3.0 - 3.0) / 2.54 * 1440)))
    if section.orientation == WD_ORIENT.LANDSCAPE:
        return twips_to_inches(landscape_twips)
    return twips_to_inches(portrait_twips)


def insert_picture(doc: Document, rule: LayoutRule, style_profile: dict) -> None:
    if rule.path is None or not rule.path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    max_width_in = resolve_image_max_width_inches(doc, style_profile)
    requested_width_in = rule.width_in or max_width_in
    final_width_in = min(requested_width_in, max_width_in)
    run.add_picture(str(rule.path), width=Inches(final_width_in))


def apply_run_font(
    paragraph,
    style_profile: dict,
    size: int | float | None = None,
    bold: bool | None = None,
) -> None:
    font_policy = resolve_font_policy(style_profile)
    for run in paragraph.runs:
        run.font.name = font_policy.get("ascii", run.font.name)
        run_rpr = run._element.get_or_add_rPr()
        set_rfonts(run_rpr, font_policy)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold


def heading_style_overrides(style_name: str, style_profile: dict) -> tuple[int | float, bool | None]:
    defaults = {"Heading 1": (18, True), "Heading 2": (16, True), "Heading 3": (14, True), "Heading 4": (12, True), "Heading 5": (12, True)}
    default_size, default_bold = defaults.get(style_name, (12, True))
    styles_cfg = style_profile.get("styles", {}) if isinstance(style_profile, dict) else {}
    cfg = styles_cfg.get(style_name, {}) if isinstance(styles_cfg, dict) else {}
    size = cfg.get("font_size_pt", default_size)
    bold = cfg.get("bold", default_bold)
    return size, bold


def render_table(doc: Document, table_def: TableDef, style_profile: dict) -> None:
    if not table_def.rows:
        return

    def normalize_col_widths(col_widths: list[int], n_cols: int, target_total: int) -> list[int]:
        if n_cols <= 0:
            return []
        if len(col_widths) != n_cols or not any(col_widths):
            authored = [1] * n_cols
        else:
            authored = [max(1, int(width)) for width in col_widths]

        authored_total = sum(authored)
        scaled: list[int] = []
        prev_target = 0
        running_authored = 0
        for idx, width in enumerate(authored, start=1):
            running_authored += width
            target_edge = round(running_authored * target_total / authored_total)
            cell_width = max(1, target_edge - prev_target)
            if idx == n_cols:
                cell_width = max(1, target_total - sum(scaled))
            scaled.append(cell_width)
            prev_target += cell_width
        return scaled

    table_cfg = style_profile.get("table_layout", {})
    n_cols = len(table_def.rows[0])
    target_total = (
        int(table_cfg.get("landscape_textbox_twips", DEFAULT_STYLE_PROFILE_DATA["table_layout"]["landscape_textbox_twips"]))
        if str(table_def.page_layout).lower() == "landscape"
        else int(table_cfg.get("portrait_textbox_twips", DEFAULT_STYLE_PROFILE_DATA["table_layout"]["portrait_textbox_twips"]))
    )
    scaled_widths = normalize_col_widths(table_def.col_widths, n_cols, target_total)

    table = doc.add_table(rows=len(table_def.rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Apply explicit borders so grid lines remain visible after front-matter merge.
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(target_total))
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    else:
        for child in list(tbl_borders):
            tbl_borders.remove(child)
    border_cfg = table_cfg.get("border", {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), str(border_cfg.get("val", "single")))
        border.set(qn("w:sz"), str(border_cfg.get("sz", "4")))
        border.set(qn("w:space"), str(border_cfg.get("space", "0")))
        border.set(qn("w:color"), str(border_cfg.get("color", "auto")))
        tbl_borders.append(border)

    for row_idx, row in enumerate(table_def.rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            normalized_value = str(value)
            cell_lines = re.split(r"<br\s*/?>|\n", normalized_value)
            if not cell_lines:
                cell_lines = [""]
            first_paragraph = cell.paragraphs[0]
            first_paragraph.clear()
            for line_idx, line in enumerate(cell_lines):
                if line_idx:
                    first_paragraph.add_run().add_break()
                first_paragraph.add_run(line)
            if col_idx < len(scaled_widths):
                cell.width = Pt(scaled_widths[col_idx] / 20)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                apply_run_font(paragraph, style_profile, size=table_def.font_size, bold=(row_idx == 0))


def chapter_heading_text(chapter: ShadowChapter) -> str:
    if chapter.chapter_id.startswith("ch"):
        return chapter_title_text(chapter.label)
    return chapter.label


def first_body_heading(chapter: ShadowChapter) -> str | None:
    if not chapter.paragraphs:
        return None
    if chapter.chapter_id.startswith("ch"):
        return chapter_title_text(chapter.label)
    return chapter.label


def resolve_table(text: str, table_defs: list[TableDef]) -> TableDef | None:
    for table_def in table_defs:
        if text.startswith(table_def.title_match):
            return table_def
    return None


def is_note_or_source(text: str) -> bool:
    return text.startswith("註：") or text.startswith("資料來源：")


def is_source_line(text: str) -> bool:
    return text.startswith("資料來源：")


def is_interview_prompt_like(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if INTERVIEW_Q_RE.match(stripped):
        return True
    if stripped.startswith(("可否", "是否", "如果", "遇到", "最後", "有沒有", "您", "判不出來", "當", "哪些")):
        return True
    if stripped.startswith(("「", "『")) and "？" in stripped:
        return True
    return False


def is_true_caption(text: str) -> bool:
    stripped = text.strip()
    if not CAPTION_RE.match(stripped):
        return False
    # Prevent long explanatory paragraphs (e.g. "表 5 ...。") from being treated as captions.
    if stripped.endswith(("。", "；", "：", "！", "？")):
        return False
    return len(stripped) <= 80


def matching_style_rule(text: str, rules: list[LayoutRule]) -> str | None:
    for rule in rules:
        if rule.kind == "style" and rule.style_name and rule.match.search(text):
            return rule.style_name
    return None


def next_nonempty_text(paragraphs: list[str], start_index: int) -> str | None:
    for idx in range(start_index, len(paragraphs)):
        stripped = visible_text(paragraphs[idx]).strip()
        if stripped:
            return stripped
    return None


def classify_body_style(paragraphs: list[str], index: int, active_item_level: int | None, prev_nonempty_text: str | None) -> str:
    next_text = next_nonempty_text(paragraphs, index + 1)
    prev_text = (prev_nonempty_text or "").strip()

    if next_text and TOP_ITEM_RE.match(next_text) and not TOP_ITEM_RE.match(prev_text):
        return "Thesis Level 4 Lead"
    if next_text and SUB_ITEM_RE.match(next_text) and TOP_ITEM_RE.match(prev_text):
        return "Thesis Level 4 Lead"
    if active_item_level == 5:
        return "Thesis Level 5 Body"
    if active_item_level == 4:
        return "Thesis Level 4 Body"
    return "Normal"


def resolve_chapter_style_rule(chapter_id: str, text: str, style_profile: dict) -> str | None:
    chapter_rules = style_profile.get("chapter_style_rules", {})
    if not isinstance(chapter_rules, dict):
        return None

    rule_cfg = chapter_rules.get(chapter_id, {})
    if not isinstance(rule_cfg, dict):
        return None

    stripped = (text or "").strip()
    if not stripped:
        return None

    exact_map = rule_cfg.get("exact_text_styles", {})
    if isinstance(exact_map, dict):
        style_name = exact_map.get(stripped)
        if isinstance(style_name, str) and style_name.strip():
            return style_name.strip()

    regex_rules = rule_cfg.get("regex_styles", [])
    if isinstance(regex_rules, list):
        for item in regex_rules:
            if not isinstance(item, dict):
                continue
            pattern = str(item.get("pattern", "")).strip()
            style_name = str(item.get("style", "")).strip()
            if not pattern or not style_name:
                continue
            try:
                if re.match(pattern, stripped):
                    return style_name
            except re.error:
                continue

    default_style = rule_cfg.get("default_style")
    if isinstance(default_style, str) and default_style.strip():
        return default_style.strip()

    return None


def classify_paragraph_style(
    paragraphs: list[str],
    index: int,
    active_item_level: int | None,
    prev_nonempty_text: str | None,
    layout_rules: list[LayoutRule] | None = None,
    legacy_style_map: dict[str, str] | None = None,
) -> tuple[str, int | None]:
    hint, stripped = strip_style_hints(paragraphs[index])
    hinted_style = style_hint_to_style_name(hint)
    if hinted_style:
        next_level = 4 if hinted_style.startswith("Thesis Level 4") else 5 if hinted_style.startswith("Thesis Level 5") else None
        return hinted_style, next_level
    if layout_rules:
        matched_style = matching_style_rule(stripped, layout_rules)
        if matched_style:
            next_level = 4 if matched_style.startswith("Thesis Level 4") else 5 if matched_style.startswith("Thesis Level 5") else None
            return matched_style, next_level
    # Structural patterns take priority over legacy style map
    if is_body_chapter_heading(stripped) or stripped == "參考文獻":
        return "Heading 1", None
    if APPENDIX_MAIN_RE.match(stripped):
        return "Appendix Heading", None
    if is_true_caption(stripped):
        sub = "圖" if re.match(r"^(附圖|圖)", stripped) else "表"
        return f"Caption {sub}", None
    if is_note_or_source(stripped):
        return "圖表註", active_item_level
    if is_interview_prompt_like(stripped):
        return classify_body_style(paragraphs, index, active_item_level, prev_nonempty_text), active_item_level
    if legacy_style_map and stripped in legacy_style_map:
        legacy_style = legacy_style_map[stripped]
        next_level = 4 if legacy_style.startswith("Thesis Level 4") else 5 if legacy_style.startswith("Thesis Level 5") else None
        return legacy_style, next_level
    if TOP_ITEM_RE.match(stripped):
        return "Thesis Level 4", 4
    if SUB_ITEM_RE.match(stripped):
        return "Thesis Level 5", 5
    if looks_like_short_heading(stripped):
        return "Heading 2", None
    return classify_body_style(paragraphs, index, active_item_level, prev_nonempty_text), active_item_level


def add_styled_paragraph(doc: Document, text: str, style_name: str) -> object:
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.add_run(text)
    return paragraph


def add_question_prompt_paragraph(doc: Document, text: str, style_name: str, style_profile: dict) -> object:
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False
    stripped = text.strip()
    match = INTERVIEW_Q_LABEL_RE.match(stripped)
    if not match:
        paragraph.add_run(text)
        apply_run_font(paragraph, style_profile, size=12, bold=False)
        return paragraph

    label, remainder = match.groups()
    label_run = paragraph.add_run(label)
    body_run = paragraph.add_run(remainder)
    apply_run_font(paragraph, style_profile, size=12, bold=None)
    label_run.font.bold = True
    body_run.font.bold = False
    return paragraph


def clear_paragraph_numbering(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


def render_chapter(
    doc: Document,
    chapter: ShadowChapter,
    table_defs: list[TableDef],
    layout_rules: list[LayoutRule],
    p_style_overrides: dict[int, str],
    p_section_break_before: dict[int, str],
    consumed_section_break_before: set[int],
    style_profile: dict,
    chapter_index: int,
) -> None:
    legacy_style_map = load_legacy_style_map()
    thesis_item_auto_numbering = is_thesis_item_numbering_enabled(style_profile)
    section_policy = resolve_section_policy(style_profile)

    def add_section_safe(orientation: str) -> None:
        add_section(
            doc,
            orientation,
            dedupe_adjacent=section_policy.get("dedupe_adjacent_section_breaks", True),
        )

    # Chapter-start section breaks are now controlled by p whitelist only.
    # Fallback policy-based chapter breaks are intentionally disabled.
    # For appendices, keep only appendix main headings (e.g., "附錄 A ...")
    # and suppress the standalone "附錄" chapter heading paragraph.
    chapter_heading_inserted = chapter.chapter_id != "appendices"
    if chapter_heading_inserted:
        chapter_heading = add_styled_paragraph(doc, chapter_heading_text(chapter), "Heading 1")
        clear_paragraph_numbering(chapter_heading)
        apply_run_font(chapter_heading, style_profile, size=18, bold=True)

    skip_first = first_body_heading(chapter)
    paragraphs = chapter.paragraphs
    if skip_first and paragraphs and visible_text(paragraphs[0].text).strip() in {
        skip_first,
        chapter.label,
        chapter_heading_text(chapter),
        chapter_title_text(chapter.label),
    }:
        paragraphs = paragraphs[1:]

    chapter_anchor_p: int | None = None
    if paragraphs:
        chapter_anchor_p = paragraphs[0].index
    if chapter_anchor_p is not None:
        chapter_start_orientation = p_section_break_before.get(chapter_anchor_p)
        if chapter_start_orientation:
            add_section_safe(chapter_start_orientation)
            consumed_section_break_before.add(chapter_anchor_p)

    landscape_open = False
    active_item_level: int | None = None
    heading1_count = 1 if chapter_heading_inserted else 0
    prev_nonempty_text: str | None = None
    previous_was_source_break = False
    for index, paragraph_record in enumerate(paragraphs):
        _, stripped = strip_style_hints(paragraph_record.text)
        if not stripped:
            continue

        # Deterministic p-level control: explicit entries in p_style_overrides.json
        # always insert a section break before this paragraph.
        forced_orientation = None
        if paragraph_record.index not in consumed_section_break_before:
            forced_orientation = p_section_break_before.get(paragraph_record.index)
        if forced_orientation:
            add_section_safe(forced_orientation)
            landscape_open = forced_orientation == "landscape"
            consumed_section_break_before.add(paragraph_record.index)
            previous_was_source_break = False

        override_style = normalize_style_alias(p_style_overrides.get(paragraph_record.index, ""))
        if override_style:
            style_name = override_style
            next_item_level = 4 if style_name.startswith("Thesis Level 4") else 5 if style_name.startswith("Thesis Level 5") else None
        else:
            # p-style architecture: unmatched paragraphs use configured default body style.
            style_name = str(style_profile.get("fallback_style", "Normal"))
            next_item_level = None
            chapter_rule_style = resolve_chapter_style_rule(chapter.chapter_id, stripped, style_profile)
            if chapter_rule_style:
                style_name = chapter_rule_style

        # Keep landscape sections open through consecutive note/source paragraphs.
        # This avoids split flows like: caption -> "註：" -> page break -> note body.
        if landscape_open and style_name != "圖表註":
            add_section_safe("portrait")
            landscape_open = False
            previous_was_source_break = False

        if style_name == "Appendix Heading":
            if section_policy.get("break_before_appendix_heading", True):
                add_section_safe("portrait")
            para = add_styled_paragraph(doc, stripped, style_name)
            clear_paragraph_numbering(para)
            set_outline_level(para, 0)
            para.paragraph_format.page_break_before = True
            apply_run_font(para, style_profile, size=18, bold=True)

            table_def = resolve_table(stripped, table_defs)
            if section_policy.get("allow_table_under_appendix_heading", True) and table_def:
                if (
                    table_def.page_layout == "landscape"
                    and section_policy.get("ensure_landscape_sections_for_wide_tables", True)
                ):
                    add_section_safe("landscape")
                    landscape_open = True
                render_table(doc, table_def, style_profile)

            active_item_level = None
            prev_nonempty_text = stripped
            continue

        if style_name in {"Caption", "Caption 圖", "Caption 表"}:
            table_def = resolve_table(stripped, table_defs)
            skip_caption_break = (
                section_policy.get("skip_caption_break_if_previous_was_source_break", True)
                and previous_was_source_break
            )
            if section_policy.get("break_before_all_captions", False) and not skip_caption_break:
                caption_orientation = "landscape" if (table_def and table_def.page_layout == "landscape") else "portrait"
                add_section_safe(caption_orientation)
                landscape_open = caption_orientation == "landscape"
            if (
                table_def
                and table_def.page_layout == "landscape"
                and section_policy.get("ensure_landscape_sections_for_wide_tables", True)
                and not landscape_open
            ):
                add_section_safe("landscape")
                landscape_open = True
            image_rule = matching_image_rule(stripped, layout_rules)
            is_figure_caption = image_rule is not None and table_def is None and stripped.startswith(("圖", "附圖"))
            if is_figure_caption:
                insert_picture(doc, image_rule, style_profile)
            para = add_styled_paragraph(doc, stripped, style_name)
            apply_run_font(para, style_profile, size=11, bold=True)
            if table_def:
                render_table(doc, table_def, style_profile)
            if image_rule and not is_figure_caption:
                insert_picture(doc, image_rule, style_profile)
            active_item_level = None
            prev_nonempty_text = stripped
            previous_was_source_break = False
            continue

        if style_name == "圖表註":
            para = add_styled_paragraph(doc, stripped, style_name)
            apply_run_font(para, style_profile, size=10.5, bold=False)
            indent = SOURCE_LINE_INDENT if is_source_line(stripped) else NOTE_LINE_INDENT
            para.paragraph_format.left_indent = indent
            para.paragraph_format.first_line_indent = Pt(-indent.pt)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.0
            if section_policy.get("break_after_source_note", False) and is_source_line(stripped):
                add_section_safe("portrait")
                landscape_open = False
                previous_was_source_break = True
            else:
                previous_was_source_break = False
            prev_nonempty_text = stripped
            continue

        if style_name == "Thesis Level 4":
            table_def = resolve_table(stripped, table_defs)
            if table_def and table_def.page_layout == "landscape":
                add_section_safe("landscape")
                landscape_open = True
            level4_text = strip_level_item_prefix(stripped, style_name) if thesis_item_auto_numbering else stripped
            if INTERVIEW_Q_LABEL_RE.match(level4_text):
                para = add_question_prompt_paragraph(doc, level4_text, style_name, style_profile)
            else:
                para = add_styled_paragraph(doc, level4_text, style_name)
                apply_run_font(para, style_profile, size=12, bold=True)
            if table_def:
                render_table(doc, table_def, style_profile)
            active_item_level = 4
            prev_nonempty_text = stripped
            previous_was_source_break = False
            continue

        if style_name == "Thesis Level 5":
            level5_text = strip_level_item_prefix(stripped, style_name) if thesis_item_auto_numbering else stripped
            if INTERVIEW_Q_LABEL_RE.match(level5_text):
                para = add_question_prompt_paragraph(doc, level5_text, style_name, style_profile)
            else:
                para = add_styled_paragraph(doc, level5_text, style_name)
                apply_run_font(para, style_profile, size=12, bold=True)
            active_item_level = 5
            prev_nonempty_text = stripped
            previous_was_source_break = False
            continue

        if style_name in {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"}:
            if style_name == "Heading 1" and heading1_count >= 1 and section_policy.get("break_before_heading_1", True):
                add_section_safe("portrait")
            if style_name == "Heading 1":
                heading1_count += 1
            table_def = resolve_table(stripped, table_defs)
            if (
                table_def
                and table_def.page_layout == "landscape"
                and section_policy.get("ensure_landscape_sections_for_wide_tables", True)
            ):
                add_section_safe("landscape")
                landscape_open = True
            para = add_styled_paragraph(doc, stripped, style_name)
            apply_page_break_rules(para, stripped, layout_rules)
            heading_size, heading_bold = heading_style_overrides(style_name, style_profile)
            apply_run_font(para, style_profile, size=heading_size, bold=heading_bold)
            if style_name == "Heading 1" and stripped == "參考文獻":
                para.paragraph_format.space_after = Pt(12)
            if table_def:
                render_table(doc, table_def, style_profile)
            active_item_level = None
            prev_nonempty_text = stripped
            previous_was_source_break = False
            continue

        active_item_level = next_item_level
        body_style = style_name
        para = add_styled_paragraph(doc, stripped.replace("\n", " "), body_style)
        apply_page_break_rules(para, stripped, layout_rules)
        apply_run_font(para, style_profile, size=12, bold=False)
        prev_nonempty_text = stripped
        previous_was_source_break = False

    if landscape_open:
        add_section_safe("portrait")


def render(
    shadow_dir: Path,
    layout_path: Path,
    template: Path,
    output: Path,
    p_style_overrides_path: Path,
    style_profile_path: Path,
) -> None:
    chapters = load_shadow_chapters(shadow_dir)
    table_defs = load_table_defs(TABLES_DIR)
    layout_rules = load_layout_rules(layout_path)
    p_style_overrides = load_p_style_overrides(p_style_overrides_path)
    p_section_break_before = load_p_section_break_before_overrides(p_style_overrides_path)
    consumed_section_break_before: set[int] = set()
    style_profile = load_style_profile(style_profile_path)

    doc = Document(template)
    ensure_v2_styles(doc, style_profile)
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
        remove_paragraph(doc.paragraphs[0])

    for chapter_index, chapter in enumerate(chapters):
        render_chapter(
            doc,
            chapter,
            table_defs,
            layout_rules,
            p_style_overrides,
            p_section_break_before,
            consumed_section_break_before,
            style_profile,
            chapter_index,
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render thesis content directly from markdown into DOCX (v2).")
    parser.add_argument("--shadow", type=Path, default=DEFAULT_SHADOW)
    parser.add_argument("--layout", type=Path, default=DEFAULT_LAYOUT)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument(
        "--p-style-overrides",
        type=Path,
        default=ROOT / "consistency" / "rules" / "p_style_overrides.json",
    )
    parser.add_argument(
        "--style-profile",
        type=Path,
        default=DEFAULT_STYLE_PROFILE,
    )
    args = parser.parse_args()
    render(
        args.shadow.resolve(),
        args.layout.resolve(),
        args.template.resolve(),
        args.output.resolve(),
        args.p_style_overrides.resolve(),
        args.style_profile.resolve(),
    )


if __name__ == "__main__":
    main()
